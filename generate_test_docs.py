import os
from docx import Document
from docx.shared import Pt


def create_test_documents():
    """
    在名为 'test_docs' 的子目录中创建几个用于测试的 Word 文档。
    """
    # 定义测试文档的内容和文件名
    # 使用 placehold.co 服务生成占位图片
    sample_docs_data = [
        {
            "filename": "测试文档_正常情况.docx",
            "title": "自由职业情况说明 - 范例一",
            "content": [
                "姓名: 张三\t\t\t\t学号: 123456789",
                "学院: 计算机学院\t\t\t学历: 硕士研究生",
                "工作内容: 作为一名自由撰稿人，主要为科技类网站撰写评测文章。",
                "学生手写签字: ![](https://placehold.co/400x80/png?text=张三的签名)",
                "学院就业负责人审核签字: ![](https://placehold.co/400x80/png?text=学院负责人签名)",
                "--- 文档结束 ---"
            ]
        },
        {
            "filename": "测试文档_无图片.docx",
            "title": "自由职业情况说明 - 范例二",
            "content": [
                "姓名: 李四\t\t\t\t学号: 987654321",
                "学院: 外国语学院\t\t\t学历: 硕士研究生",
                "工作内容: 从事兼职英语翻译工作，但本文档中没有图片链接。",
                "学生手写签字: ",
                "学院就业负责人审核签字: ",
                "--- 文档结束 ---"
            ]
        },
        {
            "filename": "测试文档_多图片混合.docx",
            "title": "自由职业情况说明 - 范例三",
            "content": [
                "姓名: 王五\t\t\t\t学号: 112233445",
                "学院: 艺术设计学院\t\t\t学历: 硕士研究生",
                "工作内容: 独立设计师，为客户提供Logo设计服务。",
                "学生手写签字: ![](https://placehold.co/400x80/0000FF/FFFFFF/png?text=王五的潇洒签名)",
                "这是一段普通的文本，不包含任何链接。",
                "学院就业负责人审核签字: ![](https://placehold.co/400x80/FF0000/FFFFFF/png?text=学院审核通过签名)",
                "学校就业部门负责人审核签字: ![](https://placehold.co/400x80/008000/FFFFFF/png?text=学校部门审核签名)",
                "--- 文档结束 ---"
            ]
        }
    ]

    # 创建用于存放测试文档的目录
    output_dir = "test_docs"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建目录: {output_dir}")

    # 循环生成每个文档
    for doc_data in sample_docs_data:
        try:
            doc = Document()

            # 添加标题
            title = doc.add_heading(doc_data["title"], level=1)
            title.alignment = 1  # 居中对齐

            # 添加内容
            for line in doc_data["content"]:
                p = doc.add_paragraph()
                # 为制表符设置一个默认字体，以防格式问题
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = Pt(12)

            # 保存文档
            save_path = os.path.join(output_dir, doc_data["filename"])
            doc.save(save_path)
            print(f"成功生成测试文档: {save_path}")

        except Exception as e:
            print(f"生成文档 {doc_data['filename']} 时出错: {e}")

    print("\n所有测试文档已生成完毕！")


if __name__ == "__main__":
    # --- 检查并提示安装依赖 ---
    try:
        import docx
    except ImportError:
        print("错误: 缺少 'python-docx' 库。")
        print("请在命令行中使用以下命令安装:\n\npip install python-docx")
        exit()

    create_test_documents()

