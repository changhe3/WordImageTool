import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Listbox, END
import threading
import re
import requests
import io
from docx import Document
from docx.shared import Inches, Emu
import os

# 尝试导入Pillow库，它是新功能的核心
try:
    from PIL import Image
except ImportError:
    # 在GUI启动前，如果缺少库，就无法继续
    # 在主程序入口处会进行更友好的提示
    pass


# --- Core Logic (核心处理逻辑已更新，以支持自定义输出路径) ---

def process_paragraph(para, log_callback, available_width=None, available_height=None):
    """
    处理单个段落，查找链接并用按比例缩放的图片替换。
    现在同时考虑宽度和高度约束。
    返回 True 表示有改动，False 表示无改动。
    """
    url_pattern = re.compile(r'(!\[.*?\]\((.*?)\))')

    if '!' not in para.text or 'http' not in para.text or not url_pattern.search(para.text):
        return False

    log_callback(f"  -> 在段落/单元格中找到链接，开始重建...")

    matches = list(url_pattern.finditer(para.text))

    segments = []
    last_end = 0
    for match in matches:
        segments.append(('text', para.text[last_end:match.start()]))
        segments.append(('image', match.group(2), match.group(1)))  # (url, original_tag)
        last_end = match.end()
    segments.append(('text', para.text[last_end:]))

    p_element = para._p
    p_element.clear_content()

    for seg_type, content, *original_tag_tuple in segments:
        if seg_type == 'text':
            if content:
                para.add_run(content)

        elif seg_type == 'image':
            url = content
            original_tag = original_tag_tuple[0]
            log_callback(f"    -> 准备处理图片链接: {url}")
            try:
                log_callback("      -> 正在下载图片...")
                response = requests.get(url, timeout=20, headers={'User-Agent': 'Mozilla/5.0'})
                response.raise_for_status()
                image_stream = io.BytesIO(response.content)
                log_callback("      -> 图片下载成功。")

                image_stream.seek(0)
                img = Image.open(image_stream)
                img_width, img_height = img.size

                if img_width == 0 or img_height == 0:
                    log_callback("      -> [警告] 图片尺寸为0，跳过处理。")
                    para.add_run(original_tag)
                    continue

                img_aspect = float(img_width) / float(img_height)
                log_callback(f"      -> 图片原始尺寸: {img_width}x{img_height}, 宽高比: {img_aspect:.2f}")

                if available_width and available_width > 0 and available_height and available_height > 0:
                    cell_width = int(available_width * 0.95)
                    cell_height = int(available_height * 0.95)
                    cell_aspect = float(cell_width) / float(cell_height)

                    if img_aspect > cell_aspect:
                        target_width = cell_width
                        target_height = int(target_width / img_aspect)
                        log_callback("      -> 图片较宽，以单元格宽度为基准缩放。")
                    else:
                        target_height = cell_height
                        target_width = int(target_height * img_aspect)
                        log_callback("      -> 图片较高，以单元格高度为基准缩放。")
                else:
                    target_width = Inches(2.0)
                    target_height = int(target_width / img_aspect)
                    log_callback("      -> 应用默认宽度。")

                log_callback(f"      -> 计算新尺寸: {target_width / 914400:.2f}\" x {target_height / 914400:.2f}\"")

                image_stream.seek(0)
                run = para.add_run()
                run.add_picture(image_stream, width=Emu(target_width), height=Emu(target_height))

            except requests.exceptions.RequestException as e:
                log_callback(f"      -> [错误] 图片下载失败: {e}")
                para.add_run(original_tag)
            except Exception as e:
                log_callback(f"      -> [错误] 插入图片时出错: {e}")
                para.add_run(original_tag)

    return True


def find_and_replace_images_in_doc(doc_path, output_dir, log_callback):
    """
    打开 .docx 文件，替换图片链接，并保存到指定的输出文件夹。
    """
    try:
        doc = Document(doc_path)
        doc_changed = False

        log_callback("--> 正在扫描文档主体段落...")
        for para in doc.paragraphs:
            if process_paragraph(para, log_callback):
                doc_changed = True

        log_callback("--> 正在扫描表格...")
        if doc.tables:
            for i, table in enumerate(doc.tables):
                log_callback(f"  -> 正在处理表格 {i + 1}...")
                for row in table.rows:
                    row_height_emu = row.height if row.height else 0
                    for cell in row.cells:
                        cell_width_emu = cell.width if cell.width else 0
                        for para in cell.paragraphs:
                            if process_paragraph(para, log_callback,
                                                 available_width=cell_width_emu,
                                                 available_height=row_height_emu):
                                doc_changed = True
        else:
            log_callback("  -> 文档中没有表格。")

        if not doc_changed:
            log_callback("  -> 未在文档中找到符合格式的图片链接或未做出任何更改。")
            return "no_images_found"

        # 使用原始文件名保存到指定的输出目录
        original_filename = os.path.basename(doc_path)
        new_doc_path = os.path.join(output_dir, original_filename)

        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

        log_callback(f"  -> 处理完成，正在保存到: {new_doc_path}")
        doc.save(new_doc_path)
        return new_doc_path

    except Exception as e:
        log_callback(f"  -> [严重错误] 处理文档时发生意外: {e}")
        import traceback
        log_callback(traceback.format_exc())
        return None


# --- Tkinter GUI Application (界面已更新) ---

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Word 图片链接批量替换工具 (自定义输出)")
        self.root.geometry("800x700")  # 增加窗口高度以容纳新控件
        self.file_paths = []
        self.output_dir = None

        self.root.configure(bg="#f0f0f0")
        font_main = ("Arial", 12)
        font_log = ("Courier New", 10)

        main_frame = tk.Frame(root, padx=20, pady=20, bg="#f0f0f0")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # --- 步骤 1: 文件选择 ---
        select_frame = tk.Frame(main_frame, bg="#f0f0f0")
        select_frame.pack(fill=tk.X, pady=(0, 10))

        self.select_button = tk.Button(select_frame, text="1. 选择Word文档 (可多选)", command=self.select_files,
                                       font=font_main, bg="#0078d4", fg="white", relief=tk.FLAT, padx=10, pady=5)
        self.select_button.pack(side=tk.LEFT, ipadx=10)

        self.clear_button = tk.Button(select_frame, text="清空列表", command=self.clear_list, font=font_main,
                                      bg="#6c757d", fg="white", relief=tk.FLAT, padx=10, pady=5)
        self.clear_button.pack(side=tk.LEFT, padx=(10, 0))

        list_frame = tk.LabelFrame(main_frame, text="待处理文件列表", font=font_main, padx=10, pady=10, bg="#f0f0f0")
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.file_listbox = Listbox(list_frame, font=("Arial", 11), relief=tk.SOLID, bd=1)
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = tk.Scrollbar(list_frame, orient="vertical", command=self.file_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill="y")
        self.file_listbox.config(yscrollcommand=scrollbar.set)

        # --- 步骤 2: 设置输出文件夹 ---
        output_frame = tk.Frame(main_frame, bg="#f0f0f0")
        output_frame.pack(fill=tk.X, pady=(10, 10))

        self.output_dir_button = tk.Button(output_frame, text="2. 设置输出文件夹", command=self.select_output_dir,
                                           font=font_main, bg="#0078d4", fg="white", relief=tk.FLAT, padx=10, pady=5)
        self.output_dir_button.pack(side=tk.LEFT, ipadx=10)

        self.output_dir_label = tk.Label(output_frame, text="尚未设置输出文件夹", font=font_main, bg="#e1e1e1",
                                         fg="#333", relief=tk.SUNKEN, anchor='w', padx=10)
        self.output_dir_label.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(10, 0))

        # --- 步骤 3: 开始处理 ---
        self.process_button = tk.Button(main_frame, text="3. 开始批量处理", command=self.start_processing_thread,
                                        font=font_main, bg="#28a745", fg="white", state=tk.DISABLED, relief=tk.FLAT,
                                        padx=10, pady=5)
        self.process_button.pack(fill=tk.X, pady=10, ipady=5)

        # --- 日志输出 ---
        log_frame = tk.LabelFrame(main_frame, text="处理日志", font=font_main, padx=10, pady=10, bg="#f0f0f0")
        log_frame.pack(fill=tk.BOTH, expand=True)

        self.log_text = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, state=tk.DISABLED, font=font_log,
                                                  relief=tk.SOLID, bd=1, height=10)
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def log(self, message):
        def _log():
            self.log_text.config(state=tk.NORMAL)
            self.log_text.insert(tk.END, message + "\n")
            self.log_text.see(tk.END)
            self.log_text.config(state=tk.DISABLED)

        self.root.after(0, _log)

    def _update_process_button_state(self):
        """根据输入和输出是否都已设置来更新处理按钮的状态"""
        if self.file_paths and self.output_dir:
            self.process_button.config(state=tk.NORMAL)
        else:
            self.process_button.config(state=tk.DISABLED)

    def select_files(self):
        paths = filedialog.askopenfilenames(
            title="请选择一个或多个Word文档",
            filetypes=[("Word Documents", "*.docx")]
        )
        if paths:
            for p in paths:
                if p not in self.file_paths:
                    self.file_paths.append(p)
            self.update_file_listbox()
            self.log(f"添加了 {len(paths)} 个文件到处理列表。当前共 {len(self.file_paths)} 个。")
        self._update_process_button_state()

    def select_output_dir(self):
        """选择输出目录"""
        path = filedialog.askdirectory(title="请选择一个文件夹用于存放处理后的文件")
        if path:
            self.output_dir = path
            self.output_dir_label.config(text=path)
            self.log(f"已设置输出文件夹: {path}")
        self._update_process_button_state()

    def update_file_listbox(self):
        self.file_listbox.delete(0, END)
        for path in self.file_paths:
            self.file_listbox.insert(END, os.path.basename(path))

    def clear_list(self):
        self.file_paths = []
        self.update_file_listbox()
        self.log("文件列表已清空。")
        self._update_process_button_state()

    def start_processing_thread(self):
        if not self.file_paths or not self.output_dir:
            messagebox.showerror("错误", "请先选择待处理的Word文档并设置输出文件夹！")
            return

        self.process_button.config(state=tk.DISABLED, text="正在处理中...")
        self.select_button.config(state=tk.DISABLED)
        self.clear_button.config(state=tk.DISABLED)
        self.output_dir_button.config(state=tk.DISABLED)

        thread = threading.Thread(target=self.process_worker, daemon=True)
        thread.start()

    def process_worker(self):
        total_files = len(self.file_paths)
        success_count = 0
        fail_count = 0
        paths_to_process = list(self.file_paths)

        for i, doc_path in enumerate(paths_to_process):
            self.log("=" * 60)
            self.log(f"开始处理文件 {i + 1}/{total_files}: {os.path.basename(doc_path)}")

            # 传递输出目录到核心函数
            new_file_path = find_and_replace_images_in_doc(doc_path, self.output_dir, self.log)

            if new_file_path and new_file_path != "no_images_found":
                success_count += 1
            elif new_file_path is None:
                fail_count += 1

        self.log("=" * 60)
        self.log(f"批量处理完成！")
        summary = f"共处理 {total_files} 个文件。\n\n成功生成: {success_count} 个\n处理失败: {fail_count} 个"
        self.log(summary.replace('\n\n', '\n'))
        messagebox.showinfo("批量处理完成", summary)

        def _reset_ui():
            self.clear_list()
            self.process_button.config(text="3. 开始批量处理")
            self.select_button.config(state=tk.NORMAL)
            self.clear_button.config(state=tk.NORMAL)
            self.output_dir_button.config(state=tk.NORMAL)
            self._update_process_button_state()

        self.root.after(0, _reset_ui)


if __name__ == "__main__":
    try:
        from PIL import Image
        import docx
        import requests
    except ImportError:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror(
            "缺少依赖库",
            "运行本程序需要'python-docx', 'requests' 和 'Pillow' 库。\n请在命令行中使用以下命令安装:\n\npip install python-docx requests Pillow"
        )
        exit()

    root = tk.Tk()
    app = App(root)
    root.mainloop()
